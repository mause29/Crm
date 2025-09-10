from docx import Document
import datetime

def create_contract(client_name, service_name):
    doc = Document()
    doc.add_heading("Contrato de Servicios", 0)
    doc.add_paragraph(f"Cliente: {client_name}")
    doc.add_paragraph(f"Servicio: {service_name}")
    doc.add_paragraph(f"Fecha: {datetime.date.today()}")
    filename = f"contract_{client_name.replace(' ', '_')}.docx"
    doc.save(filename)
    return filename
